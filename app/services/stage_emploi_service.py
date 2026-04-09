"""
Scripts pédagogiques sous `Input/Script/<parcours>/` : leçons .docx, dossier Sources/, textes améliorés.

Compat : `PROMIND7_STAGE_EMPLOI_SCRIPT_ROOT` force la racine uniquement pour le dossier
« Parcours stage & emploi ».
"""
from __future__ import annotations

import html
import os
import re
import unicodedata
from collections import Counter
from datetime import datetime
from pathlib import Path
from typing import Any, List, Optional, Tuple

from ai.document_loaders import read_docx_file, read_text_file
from ai.paths import project_root

EXPORTS_DIRNAME = "_exports"
# Par parcours : documents de référence (liens, PDF, notes…) — pas des leçons Word listées
SOURCES_DIRNAME = "Sources"
# Texte retravaillé : même arborescence relative que le .docx, fichier .md UTF-8
V2_DIRNAME = "V2"
# Second fichier Word : **même dossier** que le V1, nom ``<stem>-V2.docx`` (ex. ``…-darija-V2.docx``).
# Ancien emplacement ``…/Word_V2/…`` encore reconnu en lecture seule.
WORD_DOCX_V2_FILE_SUFFIX = "-V2"
WORD_DOCX_V2_DIRNAME = "Word_V2"
# Suffixe fréquent sur les exports rédaction (à retirer pour titre lisible / copie canonique)
SCRIPT_STEM_DARIJA_SUFFIX = "-darija"


def normalize_script_display_stem(stem: str) -> str:
    """
    Retire les marqueurs techniques en fin de nom de fichier (sans extension) :
    d’abord ``-V2``, puis ``-darija`` (insensible à la casse).

    Sert aux libellés UI, au nom proposé au téléchargement, et à la synchro vers les scripts « propres ».
    """
    s = (stem or "").strip()
    if s.endswith(WORD_DOCX_V2_FILE_SUFFIX):
        s = s[: -len(WORD_DOCX_V2_FILE_SUFFIX)]
    if s.casefold().endswith(SCRIPT_STEM_DARIJA_SUFFIX.casefold()):
        s = s[: -len(SCRIPT_STEM_DARIJA_SUFFIX)]
    return s


# Clés internes pour le filtre « Module » (UI affiche un libellé lisible)
_MODULE_ROOT = "__ROOT__"
_MODULE_EXPORTS_FLAT = "__EXPORTS_RACINE__"

_STAGE_EMPLOI_PACK = "Parcours stage & emploi"
# Phase 0 (architecture) : scripts d’accroche / L1–L5 — ancien nom ``000-Vidéos-gratuites``.
STAGE_EMPLOI_PHASE0_FOLDER = "000-Initiation"


def module_bucket_for_rel(rel_posix: str) -> str:
    """
    Dossier « module » logique pour un chemin relatif au parcours (dossier sous Input/Script).

    - Fichier à la racine du parcours → ``__ROOT__``.
    - ``_exports/<Module>/fichier`` → regroupé sous ``<Module>`` (même vue que l’original).
    - ``_exports/fichier`` seul → ``__EXPORTS_RACINE__``.
    - Sinon → premier segment de chemin (ex. ``000-Initiation``, ``00-Module 0``, …).
    """
    rel = rel_posix.replace("\\", "/").strip("/")
    parts = rel.split("/") if rel else []
    if len(parts) <= 1:
        return _MODULE_ROOT
    top = parts[0]
    if top == EXPORTS_DIRNAME:
        if len(parts) >= 3:
            return parts[1]
        return _MODULE_EXPORTS_FLAT
    return top


def is_bandes_annonce_module_bucket(bucket: str) -> bool:
    """
    Dossiers « phase 0 » (initiation) : nom commençant par ``000`` + séparateur ou fin
    (ex. ``000-Initiation``). Évite ``0001``…
    """
    b = bucket.strip()
    if not b.startswith("000"):
        return False
    if len(b) == 3:
        return True
    return not b[3].isdigit()


# Numérotation pédagogique affichée : reprise à « Leçon 1 » au début de chaque bloc
# (après Live 0, puis après Live 1, puis après Live 2 — aligné sur les jalons M2 / M5 / M7).
_LIVE_SEGMENT_BUCKET_PREFIXES: tuple[tuple[int, tuple[str, ...]], ...] = (
    (1, ("00-Module", "01-Module", "02-Module")),
    (2, ("03-Module", "04-Module", "05-Module")),
    (3, ("06-Module", "07-Module")),
)

LIVE_SEGMENT_CAPTION_FR: dict[int, str] = {
    1: "phase 1 — Le Contexte (modules 0 à 2)",
    2: "phase 2 — La Préparation (modules 3 à 5)",
    3: "phase 3 — Langues et outils numériques (modules 6 et 7)",
}


def live_lesson_segment_for_bucket(bucket: str) -> Optional[int]:
    """
    Identifiant de segment pour la renumérotation affichée (1, 2 ou 3).
    ``None`` pour les dossiers hors 00–07 (ex. ``000-…``, racine) : le libellé fichier reste brut.
    """
    b = (bucket or "").strip()
    for seg_id, prefixes in _LIVE_SEGMENT_BUCKET_PREFIXES:
        if any(b.startswith(p) for p in prefixes):
            return seg_id
    return None


_LESSON_STEM_NUM_RE = re.compile(r"(?i)^(Leçon\s+)\d+(\s*[—–-]\s*)(.*)$")


def format_lesson_display_stem(stem: str, local_n: Optional[int]) -> str:
    """
    Remplace le ``Leçon N`` du nom de fichier par la numérotation locale du segment.
    Si ``local_n`` est ``None``, retourne ``stem`` inchangé (normalisé).
    """
    raw = normalize_script_display_stem((stem or "").strip())
    if local_n is None:
        return raw
    m = _LESSON_STEM_NUM_RE.match(raw)
    if m:
        return f"{m.group(1)}{local_n}{m.group(2)}{m.group(3)}"
    return f"Leçon {local_n} — {raw}"


def canonical_lesson_number_from_stem(stem: str) -> Optional[int]:
    """Numéro « technique » du parcours extrait du stem (ex. outillage ``--lesson``)."""
    m = re.search(r"(?i)Leçon\s+(\d+)", normalize_script_display_stem(stem or ""))
    if not m:
        return None
    try:
        return int(m.group(1))
    except ValueError:
        return None


def lesson_local_number_by_rel(files: List[dict]) -> dict[str, tuple[Optional[int], Optional[int]]]:
    """
    Pour chaque ``rel_posix`` : ``(segment_id | None, leçon_affichée | None)``.
    La leçon affichée compte 1, 2, … dans l’ordre des chemins **au sein du segment**.
    """
    out: dict[str, tuple[Optional[int], Optional[int]]] = {}
    sorted_files = sorted(files, key=lambda x: x["rel_posix"].lower())
    seg_count: dict[int, int] = {}
    for f in sorted_files:
        rel = f["rel_posix"]
        seg = live_lesson_segment_for_bucket(module_bucket_for_rel(rel))
        if seg is None:
            out[rel] = (None, None)
            continue
        seg_count[seg] = seg_count.get(seg, 0) + 1
        out[rel] = (seg, seg_count[seg])
    return out


def lesson_select_display_labels(
    files: List[dict],
    *,
    local_by_rel: dict[str, tuple[Optional[int], Optional[int]]],
) -> dict[str, str]:
    """Libellés pour selectbox / sommaire ; désambiguïsation si deux stems identiques."""
    formatted: dict[str, str] = {}
    for f in files:
        stem = Path(f["name"]).stem
        _, loc = local_by_rel.get(f["rel_posix"], (None, None))
        formatted[f["rel_posix"]] = format_lesson_display_stem(stem, loc)

    counts = Counter(formatted.values())
    out: dict[str, str] = {}
    for f in files:
        rel = f["rel_posix"]
        label = formatted[rel]
        if counts[label] <= 1:
            out[rel] = label
            continue
        parent = Path(rel).parent.as_posix()
        if not parent or parent == ".":
            out[rel] = label
        else:
            out[rel] = f"{label}  ({parent})"
    return out


def format_module_bucket_label(bucket: str) -> str:
    if bucket == _MODULE_ROOT:
        return "Racine du parcours"
    if bucket == _MODULE_EXPORTS_FLAT:
        return "_exports (sans sous-dossier)"
    if is_bandes_annonce_module_bucket(bucket):
        return "Initiation"
    return bucket


def ordered_module_buckets(files: List[dict]) -> List[str]:
    """
    Liste triée des buckets présents dans ``files`` (pour selectbox).

    Les dossiers **000…** (phase 0 — initiation) sont toujours **en premier**, puis le reste par ordre
    alphabétique ; enfin racine du parcours et exports plats.
    """
    keys = {module_bucket_for_rel(f["rel_posix"]) for f in files}
    specials = {_MODULE_ROOT, _MODULE_EXPORTS_FLAT}
    normals = [k for k in keys if k not in specials]
    trailers = sorted(
        (k for k in normals if is_bandes_annonce_module_bucket(k)),
        key=lambda s: s.lower(),
    )
    others = sorted(
        (k for k in normals if not is_bandes_annonce_module_bucket(k)),
        key=lambda s: s.lower(),
    )
    ordered_normals = trailers + others
    tail: List[str] = []
    if _MODULE_ROOT in keys:
        tail.append(_MODULE_ROOT)
    if _MODULE_EXPORTS_FLAT in keys:
        tail.append(_MODULE_EXPORTS_FLAT)
    return ordered_normals + tail


def get_script_library_root() -> Path:
    """Racine `Input/Script` (ou `PROMIND7_SCRIPT_LIBRARY_ROOT`)."""
    env = os.getenv("PROMIND7_SCRIPT_LIBRARY_ROOT", "").strip()
    if env:
        return Path(env)
    return project_root() / "Input" / "Script"


def consulting_stage_emploi_scripts_dir() -> Path:
    """
    Emplacement conventionnel hors dépôt Streamlit : ``01-Promind7/02-Contenus/01-Scripts/…``.

    ``project_root()`` = dossier ``01-developpement IA`` → parents ``03-Streamlit`` puis ``01-Promind7``.
    """
    return project_root().parent.parent / "02-Contenus" / "01-Scripts" / _STAGE_EMPLOI_PACK


def _stage_emploi_pack_from_consulting_layout() -> Optional[Path]:
    """Si ce dossier existe, utilisé par défaut pour Stage & emploi (sans variable d’environnement)."""
    try:
        p = consulting_stage_emploi_scripts_dir()
    except (OSError, ValueError):
        return None
    return p if p.is_dir() else None


def get_pack_root(pack_folder_name: str) -> Path:
    """Racine d’un parcours (sous-dossier de la bibliothèque de scripts)."""
    if pack_folder_name == _STAGE_EMPLOI_PACK:
        alt = os.getenv("PROMIND7_STAGE_EMPLOI_SCRIPT_ROOT", "").strip()
        if alt:
            return Path(alt)
        layout = _stage_emploi_pack_from_consulting_layout()
        if layout is not None:
            return layout
    return get_script_library_root() / pack_folder_name


def list_parcours_folder_names() -> List[str]:
    """Noms des dossiers parcours sous la bibliothèque (triés)."""
    root = get_script_library_root()
    if not root.is_dir():
        return []
    return sorted(
        p.name for p in root.iterdir() if p.is_dir() and not p.name.startswith(".")
    )


list_pack_folder_names = list_parcours_folder_names


def get_stage_emploi_root() -> Path:
    """Rétrocompat : racine du parcours Stage & emploi."""
    return get_pack_root(_STAGE_EMPLOI_PACK)


# Consolidé veille YouTube (transcripts + idées) — généré par ``build_veille_transcripts_consolidated.py``
CONSOLIDATED_VEILLE_VIDEOS_FILENAME = "CONSOLIDE_transcripts_idees_proM7.md"


def consolidated_veille_videos_path(pack_root: Path) -> Path:
    """Chemin du fichier unique regroupant toutes les transcriptions / idées (Markdown)."""
    return pack_root / "Sources" / "videos" / CONSOLIDATED_VEILLE_VIDEOS_FILENAME


# --- Sommaire UI Streamlit : grille officielle lives + tests (hors Word ; calendrier = plateforme)
#
# Séquence cible (rail « Tous les modules ») : **dossier 000-Initiation** (phase 0, contenu Word) →
# **test avant parcours** → **Live 0 (ouverture)** → 00…07 → Live 1–3 / tests de module → test après parcours.
# Réf. feuille de route produit (modules 0–7, leçons 1–62).

STAGE_EMPLOI_SOMMAIRE_OPENING_LIVE: dict[str, Any] = {
    "title": "Live 0 — ouverture",
    "bullets": [
        "Session de **lancement** du parcours Stage & emploi : objectifs, rythme (vidéos / **lives**), posture apprenant·e et rôle de **proM7**.",
        "**Créneaux** et lien de connexion : plateforme / animation.",
    ],
}

STAGE_EMPLOI_SOMMAIRE_TEST_BEFORE_PARCOURS: dict[str, Any] = {
    "title": "Test de connaissance — avant parcours",
    "bullets": [
        "**Évaluation initiale** (plateforme) avant d’entrer dans les modules — repère de départ, sans jugement.",
    ],
}

STAGE_EMPLOI_SOMMAIRE_TEST_AFTER_PARCOURS: dict[str, Any] = {
    "title": "Test de connaissance — après parcours",
    "bullets": [
        "**Évaluation finale** (plateforme) en clôture du parcours — bilan des acquis par rapport au trajet 00–07.",
    ],
}

_STAGES_EMPLOI_MODULE_TEST_TITLES: tuple[tuple[str, str], ...] = (
    ("00-Module", "Test — Module 0"),
    ("01-Module", "Test — Module 1"),
    ("02-Module", "Test — Module 2"),
    ("03-Module", "Test — Module 3"),
    ("04-Module", "Test — Module 4"),
    ("05-Module", "Test — Module 5"),
    ("06-Module", "Test — Module 6"),
    ("07-Module", "Test — Module 7"),
)


def sommaire_module_end_test_block(bucket_key: str) -> dict[str, Any] | None:
    """Test de fin de module (affiché après le bloc leçons du dossier-module 00–07)."""
    for substr, title in _STAGES_EMPLOI_MODULE_TEST_TITLES:
        if substr in bucket_key:
            return {
                "title": title,
                "bullets": [
                    "**Contrôle / bilan** de fin de module sur la plateforme (contenu couvert par les leçons du dossier)."
                ],
            }
    return None


STAGE_EMPLOI_SOMMAIRE_LIVES_AFTER_BUCKET: list[dict[str, Any]] = [
    {
        "bucket_substr": "02-Module",
        "title": "Live — Phase 1 — Le Contexte",
        "bullets": [
            "Après **modules 0 à 2** (mindset, économie & marché, entreprise).",
            "🔴 **Le Contexte** — نلخّصو، نجوّبو على الأسئلة، ونثبّتو هاد الأساسيات مع بعضنا (aligné sur le plan de leçon M0).",
        ],
    },
    {
        "bucket_substr": "05-Module",
        "title": "Live — Phase 2 — La Préparation",
        "bullets": [
            "Après **modules 3 à 5** (comportement, stage, emploi).",
            "🔴 **La Préparation** — نرصّو الأدوات والعادات : تقويم، أولويات، طريقة تخدم بيها على راسك.",
        ],
    },
    {
        "bucket_substr": "07-Module",
        "title": "Live — Phase 3 — Langues et outils numériques",
        "bullets": [
            "Après **modules 6 et 7** (outils numériques, langues).",
            "Session **Langues et outils numériques** : pratique, questions, passage concret **FR/ANG** et culture **digital / IA**.",
        ],
    },
]


def sommaire_live_block_for_bucket(bucket_key: str) -> dict[str, Any] | None:
    """Live global à afficher **après** le test de fin de module du dossier concerné (M2, M5, M7)."""
    for block in STAGE_EMPLOI_SOMMAIRE_LIVES_AFTER_BUCKET:
        if block["bucket_substr"] in bucket_key:
            return block
    return None


def _reserved_path_segments() -> set[str]:
    return {
        SOURCES_DIRNAME.casefold(),
        V2_DIRNAME.casefold(),
        WORD_DOCX_V2_DIRNAME.casefold(),
        "textes_ameliores",  # ancien emplacement (ne pas lister comme leçon)
        EXPORTS_DIRNAME.casefold(),
    }


def _rel_crosses_reserved_folder(rel_posix: str) -> bool:
    parts = rel_posix.replace("\\", "/").split("/")
    reserved = _reserved_path_segments()
    return any(p.casefold() in reserved for p in parts)


def get_sources_dir(pack_root: Path) -> Path:
    return Path(pack_root) / SOURCES_DIRNAME


def list_sources_files(pack_root: Path) -> List[dict]:
    """Fichiers sous ``<pack>/Sources/`` (tous types), triés par chemin relatif."""
    d = get_sources_dir(pack_root)
    if not d.is_dir():
        return []
    items: List[dict] = []
    for p in d.rglob("*"):
        if not p.is_file():
            continue
        rel = p.relative_to(d)
        items.append(
            {
                "rel_posix": rel.as_posix(),
                "path": p,
                "name": p.name,
            }
        )
    items.sort(key=lambda x: x["rel_posix"].lower())
    return items


def is_word_docx_v2_filename(name: str) -> bool:
    """Vrai si le fichier est un Word « secondaire » listé à part (suffixe ``-V2`` avant ``.docx``)."""
    p = Path(name)
    if p.suffix.lower() != ".docx":
        return False
    return p.stem.endswith(WORD_DOCX_V2_FILE_SUFFIX)


def canonical_word_docx_v2_path(pack_root: Path, docx_rel_posix: str) -> Path:
    """
    Cible d’écriture du Word V2 : **même dossier** que le V1, nom ``{stem}-V2.docx``.
    """
    v1 = Path(pack_root) / docx_rel_posix.replace("/", os.sep)
    return v1.parent / f"{v1.stem}{WORD_DOCX_V2_FILE_SUFFIX}{v1.suffix}"


def _legacy_word_v2_folder_path(pack_root: Path, docx_rel_posix: str) -> Path:
    """Ancienne convention : ``<module>/Word_V2/<suite du chemin>``."""
    rel = docx_rel_posix.replace("\\", "/").strip("/")
    p = Path(rel)
    parts = p.parts
    root = Path(pack_root)
    if len(parts) < 2:
        return root / WORD_DOCX_V2_DIRNAME / p
    top = parts[0]
    tail = Path(*parts[1:])
    return root / top / WORD_DOCX_V2_DIRNAME / tail


def locate_word_docx_v2_path(pack_root: Path, docx_rel_posix: str) -> Optional[Path]:
    """
    Retourne le chemin du fichier V2 **s’il existe** : d’abord le fichier ``-V2`` dans le même
    dossier que le V1, sinon l’emplacement legacy ``Word_V2/``.
    """
    cand = canonical_word_docx_v2_path(pack_root, docx_rel_posix)
    if cand.is_file():
        return cand
    leg = _legacy_word_v2_folder_path(pack_root, docx_rel_posix)
    if leg.is_file():
        return leg
    return None


def resolve_lesson_script_docx_path(pack_root: Path, docx_rel_posix: str) -> Optional[Path]:
    """
    Fichier Word à afficher pour une entrée « leçon » (chemin relatif vers un .docx **sans** suffixe
    ``-V2`` dans le nom, tel que retourné par ``list_lesson_docx_files``) :

    1. Jumeau ``-V2`` (convention historique : V1 listé, contenu lu depuis ``…-V2.docx``) ;
    2. Sinon le fichier listé lui-même s’il existe (copie canonique unique, sans mention V2 sur disque).
    """
    v2 = locate_word_docx_v2_path(pack_root, docx_rel_posix)
    if v2 is not None:
        return v2
    direct = Path(pack_root) / docx_rel_posix.replace("/", os.sep)
    if direct.is_file():
        return direct
    return None


def word_v2_docx_path(pack_root: Path, docx_rel_posix: str) -> Path:
    """
    Compat : préférer ``canonical_word_docx_v2_path`` (chemin d’écriture V2).
    """
    return canonical_word_docx_v2_path(pack_root, docx_rel_posix)


def v2_md_path(pack_root: Path, docx_rel_posix: str) -> Path:
    """
    Fichier **V2** : ``V2/<même dossiers que le .docx>/Nom.md`` (UTF-8).
    """
    rel = Path(docx_rel_posix)
    stem = rel.stem
    parent = rel.parent
    base = Path(pack_root) / V2_DIRNAME
    if parent == Path(".") or str(parent) == ".":
        return base / f"{stem}.md"
    return base / parent / f"{stem}.md"


def load_v2_text(pack_root: Path, docx_rel_posix: str) -> Tuple[Optional[str], Optional[str]]:
    """Lit le V2 si le fichier existe ; sinon (None, None). Erreur → (None, msg)."""
    p = v2_md_path(pack_root, docx_rel_posix)
    if not p.is_file():
        return None, None
    try:
        return p.read_text(encoding="utf-8"), None
    except Exception as exc:  # noqa: BLE001
        return None, str(exc)


def save_v2_text(
    pack_root: Path, docx_rel_posix: str, text: str
) -> Tuple[Optional[Path], Optional[str]]:
    """Enregistre / écrase le .md V2 (UTF-8)."""
    root = Path(pack_root)
    if not root.is_dir():
        return None, "Dossier parcours introuvable."
    raw = text.strip()
    if not raw:
        return None, "Texte vide."
    src = root / docx_rel_posix.replace("/", os.sep)
    if not src.is_file():
        return None, "Fichier Word source introuvable."
    out = v2_md_path(root, docx_rel_posix)
    try:
        out.parent.mkdir(parents=True, exist_ok=True)
        out.write_text(raw, encoding="utf-8")
        return out, None
    except Exception as exc:  # noqa: BLE001
        return None, str(exc)


def list_lesson_docx_files(pack_root: Path) -> List[dict]:
    """
    Leçons : uniquement ``.docx`` hors ``Sources/``, ``V2/``, ``_exports/``, ``Word_V2/``, etc.
    Les fichiers ``*-V2.docx`` (second Word dans le même dossier) ne sont pas des entrées séparées.
    """
    items = list_script_files(pack_root, include_exports=False)
    return [
        it
        for it in items
        if it["ext"] == ".docx"
        and not _rel_crosses_reserved_folder(it["rel_posix"])
        and not is_word_docx_v2_filename(it["name"])
    ]


# --- Raccourci Architecture → onglet Script (correspondance libellé « Ln — … » ↔ .docx) ------------

_ARCH_L_LINE_RE = re.compile(r"(?i)^L\s*(\d+)\s*[—–-]\s*(.+)$")


def is_architecture_lesson_line(line: str) -> bool:
    """Vrai si la ligne du schéma est une leçon « Ln — … » (pas un test ni un live)."""
    return _ARCH_L_LINE_RE.match((line or "").strip()) is not None
_STEM_LECON_TAIL_RE = re.compile(
    r"(?i)^(?:\d+\s*[—–-]\s*)?(?:Leçon\s+\d+\s*[—–-]\s*)(.+)$"
)


def _norm_arch_title(s: str) -> str:
    t = (s or "").strip().lower()
    t = "".join(
        c for c in unicodedata.normalize("NFD", t) if unicodedata.category(c) != "Mn"
    )
    t = t.replace("—", "-").replace("–", "-")
    t = re.sub(r"\s+", " ", t)
    return t


def stem_title_tail_for_architecture_match(stem: str) -> str:
    """Partie titre après « Leçon n — » sur un stem de fichier (préfixe numérique optionnel)."""
    s = normalize_script_display_stem((stem or "").strip())
    m = _STEM_LECON_TAIL_RE.match(s)
    return (m.group(1).strip() if m else "")


def pick_phase0_bucket_from_files(files: List[dict]) -> Optional[str]:
    keys = {module_bucket_for_rel(f["rel_posix"]) for f in files}
    cands = [k for k in keys if is_bandes_annonce_module_bucket(k)]
    if not cands:
        return None
    cands.sort(key=lambda x: x.lower())
    return cands[0]


def pick_module_bucket_from_files(files: List[dict], module_index: int) -> Optional[str]:
    """``module_index`` : 0..7 pour « Module 0 » … « Module 7 »."""
    prefix = f"{int(module_index):02d}-Module"
    keys = {module_bucket_for_rel(f["rel_posix"]) for f in files}
    matches = sorted((k for k in keys if k.startswith(prefix)), key=lambda s: s.lower())
    return matches[0] if matches else None


def module_index_from_architecture_module_title(module_title: str) -> Optional[int]:
    m = re.match(r"(?i)^Module\s+(\d+)\s", (module_title or "").strip())
    if not m:
        return None
    return int(m.group(1))


def _lesson_docx_entries_in_bucket(files: List[dict], bucket: str) -> List[dict]:
    out = [f for f in files if module_bucket_for_rel(f["rel_posix"]) == bucket]
    out.sort(key=lambda x: x["rel_posix"].lower())
    return out


def _looks_like_test_docx_stem(stem: str) -> bool:
    s = normalize_script_display_stem((stem or "").strip())
    return bool(re.search(r"(?i)\btest\b", s))


def resolve_architecture_lesson_line_to_rel_posix(
    pack_root: Path,
    *,
    architecture_line: str,
    bucket: Optional[str],
) -> Optional[str]:
    """
    Retourne ``rel_posix`` du .docx pour une ligne « Ln — Titre » du schéma architecture.

    Correspondance : d’abord **titre** (fuzzy, insensible aux accents / tirets), sinon **position** Ln
    parmi les .docx du dossier qui ne ressemblent pas à un fichier « Test … ».
    """
    m = _ARCH_L_LINE_RE.match((architecture_line or "").strip())
    if not m:
        return None
    local_i = int(m.group(1))
    arch_title = m.group(2).strip()
    if not bucket or not Path(pack_root).is_dir():
        return None
    all_f = list_lesson_docx_files(pack_root)
    cands = _lesson_docx_entries_in_bucket(all_f, bucket)
    lessonish = [f for f in cands if not _looks_like_test_docx_stem(Path(f["name"]).stem)]
    arch_norm = _norm_arch_title(arch_title)
    best_rel: Optional[str] = None
    best_score = 0
    for f in lessonish:
        tail = stem_title_tail_for_architecture_match(Path(f["name"]).stem)
        if not tail:
            continue
        fn_norm = _norm_arch_title(tail)
        if arch_norm == fn_norm:
            return f["rel_posix"]
        if arch_norm in fn_norm or fn_norm in arch_norm:
            score = min(len(arch_norm), len(fn_norm))
            if score > best_score:
                best_score = score
                best_rel = f["rel_posix"]
    if best_rel:
        return best_rel
    if 1 <= local_i <= len(lessonish):
        return lessonish[local_i - 1]["rel_posix"]
    return None


def stage_emploi_script_jump_session_key(key_slug: str) -> str:
    """Clé ``st.session_state`` : cible un ``rel_posix`` à ouvrir dans l’onglet Script."""
    return f"pm7_se_script_jump_{key_slug}"


def _is_under_exports(path: Path, root: Path) -> bool:
    try:
        rel = path.relative_to(root)
        return rel.parts and rel.parts[0] == EXPORTS_DIRNAME
    except ValueError:
        return False


def list_script_files(pack_root: Path, *, include_exports: bool = False) -> List[dict]:
    """
    Liste les .docx, .md, .txt sous `pack_root`, triés par chemin relatif.
    """
    root = Path(pack_root)
    if not root.is_dir():
        return []

    items: List[dict] = []
    for p in root.rglob("*"):
        if not p.is_file():
            continue
        if not include_exports and _is_under_exports(p, root):
            continue
        suf = p.suffix.lower()
        if suf not in {".docx", ".md", ".txt"}:
            continue
        rel = p.relative_to(root)
        items.append(
            {
                "rel_posix": rel.as_posix(),
                "path": p,
                "name": p.name,
                "ext": suf,
            }
        )
    items.sort(key=lambda x: x["rel_posix"].lower())
    return items


_OBJECTIF_HEADING = "الهدف"
# Titre de séquence proM7 (bleu dans les .docx générés par les format_*)
_SEQ_TITLE_LINE_RE = re.compile(r"(?i)^Séquence\s+\d+\s*—\s*.+")
# Évite « Séquence 1 — 1. … » comme dans format_module00_scripts.py
_SEQ_TITLE_DEDUP_RE = re.compile(
    r"(?i)Séquence\s+(\d+)\s*—\s*\1\s*(?:\.\s*|\s+\.\s+)(.+)",
)


def normalize_sequence_heading_line(title: str) -> str:
    """Affichage homogène des titres « Séquence N — … »."""
    t = (title or "").strip()
    if not t:
        return t
    m = _SEQ_TITLE_DEDUP_RE.match(t)
    if m:
        rest = m.group(2).strip()
        return f"Séquence {m.group(1)} — {rest}"
    return t


# --- Affichage HTML Sommaire (texte mixte FR/Latin ↔ arabe, BiDi) -----------------

_SEQ_DISPLAY_PREFIX_RE = re.compile(r"(?i)^(Séquence\s+\d+\s*—\s*)")


def _is_arabic_letter(ch: str) -> bool:
    if len(ch) != 1:
        return False
    o = ord(ch)
    return (
        0x0600 <= o <= 0x06FF
        or 0x0750 <= o <= 0x077F
        or 0x08A0 <= o <= 0x08FF
        or 0xFB50 <= o <= 0xFDFF
        or 0xFE70 <= o <= 0xFEFF
    )


def _is_arabic_digit(ch: str) -> bool:
    if len(ch) != 1:
        return False
    o = ord(ch)
    return 0x0660 <= o <= 0x0669 or 0x06F0 <= o <= 0x06F9


def _is_latin_letter(ch: str) -> bool:
    if len(ch) != 1:
        return False
    o = ord(ch)
    if ch.isascii() and ch.isalpha():
        return True
    # Français et extensions latines usuelles
    if 0x00C0 <= o <= 0x02AF:
        return True
    return False


def _strong_bidi_dir(ch: str) -> int | None:
    """-1 = RTL (arabe), 1 = LTR (latin), None = neutre (ponctuation, espaces, chiffres ASCII).

    Les chiffres 0–9 sont **neutres** : ils restent accrochés au mot arabe ou latin voisin,
    ce qui évite qu’un « 2030 » isolé en LTR casse l’ordre d’un titre entièrement arabe
    dans le Sommaire Streamlit.
    """
    if _is_arabic_letter(ch) or _is_arabic_digit(ch):
        return -1
    if _is_latin_letter(ch):
        return 1
    return None


def line_has_arabic_script(s: str) -> bool:
    return any(_is_arabic_letter(c) for c in s)


def line_has_latin_script(s: str) -> bool:
    """Latin « fort » (lettres + chiffres occidentaux), pour distinguer texte mixte du pur arabe."""
    for c in s:
        d = _strong_bidi_dir(c)
        if d == 1:
            return True
    return False


def mixed_text_to_bidi_spans_html(s: str) -> str:
    """
    Découpe une ligne en segments forts LTR / RTL (HTML échappé par segment).
    Les caractères neutres restent accrochés au segment courant ; une file initiale
    neutre est fusionnée avec le premier segment fort.
    """
    if not s:
        return ""
    segments: list[tuple[int, str]] = []
    pending_neutral = ""

    def append_segment(d: int, text: str) -> None:
        if not text:
            return
        if segments and segments[-1][0] == d:
            segments[-1] = (d, segments[-1][1] + text)
        else:
            segments.append((d, text))

    i = 0
    n = len(s)
    while i < n:
        ch = s[i]
        sd = _strong_bidi_dir(ch)
        if sd is None:
            pending_neutral += ch
            i += 1
            continue
        d = sd
        chunk = pending_neutral + ch
        pending_neutral = ""
        i += 1
        while i < n:
            ch2 = s[i]
            sd2 = _strong_bidi_dir(ch2)
            if sd2 is None:
                chunk += ch2
                i += 1
                continue
            if sd2 != d:
                break
            chunk += ch2
            i += 1
        append_segment(d, chunk)

    if pending_neutral:
        if segments:
            ld, lt = segments[-1]
            segments[-1] = (ld, lt + pending_neutral)
        else:
            append_segment(1, pending_neutral)

    parts: list[str] = []
    for d, text in segments:
        dir_attr = "rtl" if d < 0 else "ltr"
        parts.append(
            f'<span dir="{dir_attr}" style="unicode-bidi:isolate;">{html.escape(text)}</span>'
        )
    return "".join(parts)


def format_mixed_line_html(
    line: str,
    *,
    pre_wrap: bool = False,
    margin_em: str = "0.25em 0",
) -> str:
    """Un paragraphe HTML pour une ligne potentiellement mixte arabe / latin."""
    raw = line or ""
    if not raw.strip():
        return ""
    has_ar = line_has_arabic_script(raw)
    has_lat = line_has_latin_script(raw)
    ws = "white-space:pre-wrap; " if pre_wrap else ""
    if not has_ar:
        return (
            f'<p dir="ltr" style="text-align:left; {ws}margin:{margin_em};">'
            f"{html.escape(raw)}</p>"
        )
    inner = mixed_text_to_bidi_spans_html(raw)
    # Pur arabe : bloc RTL aligné à droite (confort de lecture).
    # Mixte FR/Latin + arabe : conteneur LTR comme pour les titres de séquences.
    if has_lat:
        return (
            f'<p dir="ltr" style="text-align:left; {ws}unicode-bidi:plaintext; margin:{margin_em};">'
            f"{inner}</p>"
        )
    return (
        f'<p dir="rtl" style="text-align:right; {ws}unicode-bidi:plaintext; margin:{margin_em};" '
        f'lang="ar">{inner}</p>'
    )


_OBJECTIF_AR_SEMICOLON_SPLIT_RE = re.compile(r"\s*؛\s*")


def _sommaire_one_chunk_bidi(chunk: str) -> str:
    """Un fragment sans ؛ ; découpe éventuelle sur « : » entre deux blocs arabes (évite BiDi faux)."""
    c = (chunk or "").strip()
    if not c:
        return ""
    if ":" in c and line_has_arabic_script(c):
        left, _, right = c.partition(":")
        ls, rs = left.strip(), right.strip()
        if (
            ls
            and rs
            and line_has_arabic_script(ls)
            and line_has_arabic_script(rs)
            and not re.search(r"[A-Za-z]{2,}", ls)
        ):
            colon_sep = '<span dir="ltr" style="unicode-bidi:isolate;">: </span>'
            return (
                mixed_text_to_bidi_spans_html(ls)
                + colon_sep
                + mixed_text_to_bidi_spans_html(rs)
            )
    return mixed_text_to_bidi_spans_html(c)


def sommaire_mixed_paragraph_inner(text: str) -> str:
    """
    Chaîne HTML (spans) pour une ligne d’objectif ou le corps d’un titre de séquence dans le Sommaire.
    Découpe sur ؛ puis applique _sommaire_one_chunk_bidi à chaque morceau.
    """
    raw = (text or "").strip()
    if not raw:
        return ""
    chunks_ar = [c.strip() for c in _OBJECTIF_AR_SEMICOLON_SPLIT_RE.split(raw) if c.strip()]
    parts = [_sommaire_one_chunk_bidi(x) for x in chunks_ar]
    if len(parts) > 1:
        sep = '<span dir="ltr" style="unicode-bidi:isolate;"> ؛ </span>'
        return sep.join(parts)
    return parts[0] if parts else ""


def format_sommaire_objectif_html(text: str) -> str:
    """
    Rendu Sommaire du bloc « Objectif » : stabilise le BiDi (؛, paquets latin au milieu, etc.).

    Sans découpage, une ligne du type « arabe؛ Latin؛ arabe » peut paraître tronquée ou réordonnée.
    Chaque paragraphe Word (ligne ``\\n``) produit un ``<p>`` distinct.
    """
    raw = (text or "").strip()
    if not raw:
        return ""
    blocks: list[str] = []
    for para in raw.split("\n"):
        ln = para.strip()
        if not ln:
            continue
        inner = sommaire_mixed_paragraph_inner(ln)
        if not inner:
            continue
        blocks.append(
            '<p dir="ltr" style="text-align:left; margin:0.35em 0; unicode-bidi:plaintext;">'
            f"{inner}</p>"
        )
    return "".join(blocks)


def _sequence_body_arabic_gloss(body: str) -> str | None:
    """
    Glose arabe pour les sous-titres de séquence **uniquement en latin** (ex. synthèse / micro-action).
    Les titres qui contiennent déjà de l’arabe ne sont pas modifiés.
    """
    b = (body or "").strip()
    if not b or line_has_arabic_script(b):
        return None

    if re.fullmatch(r"(?i)Micro-action", b):
        return "ميكرو-أكشن (خطوة عملية صغيرة)"

    m = re.fullmatch(
        r"(?i)Synthèse,\s*pièges\s+fréquents,\s*pont vers la leçon\s+(\d+)",
        b,
    )
    if m:
        return f"ملخص، أغلاط شائعة، جسر نحو الدرس {m.group(1)}"

    m = re.fullmatch(
        r"(?i)Synthèse,\s*pièges,\s*pont vers la leçon\s+(\d+)\s*\(PFE\)",
        b,
    )
    if m:
        return f"ملخص، أغلاط، جسر نحو الدرس {m.group(1)} (PFE)"

    m = re.fullmatch(
        r"(?i)Synthèse,\s*pièges,\s*pont vers la leçon\s+(\d+)",
        b,
    )
    if m:
        return f"ملخص، أغلاط، جسر نحو الدرس {m.group(1)}"

    m = re.fullmatch(
        r"(?i)Synthèse,\s*pièges,\s*pont vers le Module\s+(\d+)",
        b,
    )
    if m:
        return f"ملخص، أغلاط، جسر نحو الموديل {m.group(1)}"

    m = re.fullmatch(
        r"(?i)Synthèse,\s*pièges,\s*pont vers\s+(\d+)\s*[–-]\s*(\d+)",
        b,
    )
    if m:
        return f"ملخص، أغلاط، جسر نحو الدروس {m.group(1)}–{m.group(2)}"

    m = re.fullmatch(
        r"(?i)Synthèse,\s*pièges,\s*clôture du module\s+(\d+)",
        b,
    )
    if m:
        return f"ملخص، أغلاط، ختام الموديل {m.group(1)}"

    return None


def format_sommaire_sequence_html_item(item_index: int, title: str) -> str:
    """
    Ligne numérotée d’une séquence : préfixe « N. Séquence M — » en LTR isolé,
    reste segmenté LTR/RTL pour un rendu correct (ex. proM7, Talent, CV au milieu du darija).

    Le paragraphe est **toujours** en ``dir=ltr`` / aligné à gauche : ainsi toutes les lignes
    (français pur ou mixte) partent du même bord, et les blocs arabes restent en ``isolate``
    sans inverser tout le titre comme avec un ``<p dir="rtl">``.

    Pour les sous-titres **100 % latins** types « Synthèse, pièges… » / « Micro-action »,
    une **glose arabe** (MSA courte) est ajoutée après un tiret cadratin — uniquement dans le
    Sommaire UI, sans modifier les .docx.
    """
    t = (title or "").strip()
    iprefix = f"{item_index}. "
    m = _SEQ_DISPLAY_PREFIX_RE.match(t)
    if m:
        seq_prefix = m.group(1)
        body = t[m.end() :]
        ltr_leading = iprefix + seq_prefix
        leading_html = (
            f'<span dir="ltr" style="unicode-bidi:isolate;">{html.escape(ltr_leading)}</span>'
        )
        body_html = sommaire_mixed_paragraph_inner(body) if body else ""
        gloss = _sequence_body_arabic_gloss(body)
        gloss_html = ""
        if gloss:
            gloss_html = (
                '<span dir="ltr" style="unicode-bidi:isolate;"> — </span>'
                f'<span dir="rtl" style="unicode-bidi:isolate;">{html.escape(gloss)}</span>'
            )
        inner = leading_html + body_html + gloss_html
    else:
        inner = (
            f'<span dir="ltr" style="unicode-bidi:isolate;">{html.escape(iprefix)}</span>'
            f"{sommaire_mixed_paragraph_inner(t)}"
        )

    return (
        '<p dir="ltr" style="text-align:left; margin:0.35em 0; unicode-bidi:plaintext;">'
        f"{inner}</p>"
    )


def extract_lesson_sommaire_from_docx(docx_path: Path) -> dict[str, Any]:
    """
    Extrait le squelette d’une leçon Word au format proM7 : cadre (lignes avant الهدف),
    texte d’objectif (après «الهدف»), titres des séquences (paragraphes « Séquence N — … »).

    S’appuie sur la structure écrite par ``format_module*_scripts.py`` ; les .docx plus
    anciens ou atypiques peuvent retourner ``error`` non vide.
    """
    p = Path(docx_path)
    out: dict[str, Any] = {
        "cadre_lines": [],
        "objectif_text": "",
        "sequence_titles": [],
        "error": None,
    }
    if not p.is_file():
        out["error"] = "Fichier introuvable."
        return out
    try:
        import docx as python_docx  # python-docx
    except ImportError:
        out["error"] = "python-docx non installé."
        return out
    try:
        document = python_docx.Document(str(p))
    except Exception as exc:  # noqa: BLE001
        out["error"] = str(exc)
        return out

    paras = [(para.text or "").strip() for para in document.paragraphs if (para.text or "").strip()]
    idx_obj = None
    for i, text in enumerate(paras):
        if text == _OBJECTIF_HEADING:
            idx_obj = i
            break
    if idx_obj is None:
        out["error"] = "Titre «الهدف» introuvable (structure non reconnue)."
        for text in paras:
            if _SEQ_TITLE_LINE_RE.match(text):
                out["sequence_titles"].append(normalize_sequence_heading_line(text))
        return out

    for j in range(idx_obj):
        line = paras[j]
        if _SEQ_TITLE_LINE_RE.match(line):
            continue
        out["cadre_lines"].append(line)

    k = idx_obj + 1
    obj_parts: list[str] = []
    while k < len(paras) and not _SEQ_TITLE_LINE_RE.match(paras[k]):
        obj_parts.append(paras[k])
        k += 1
    out["objectif_text"] = "\n".join(obj_parts).strip()

    while k < len(paras):
        line = paras[k]
        if _SEQ_TITLE_LINE_RE.match(line):
            out["sequence_titles"].append(normalize_sequence_heading_line(line))
        k += 1
    return out


def load_script_text(path: Path, max_chars: int = 200_000) -> tuple[str, Optional[str]]:
    """Charge le texte d'un script supporté."""
    p = Path(path)
    if not p.is_file():
        return "", "Fichier introuvable."
    suf = p.suffix.lower()
    try:
        if suf in {".md", ".txt"}:
            t = read_text_file(p, max_chars=max_chars)
            return t, None if t else "Fichier vide."
        if suf == ".docx":
            t = read_docx_file(p, max_chars=max_chars)
            return t, None if t else "Lecture vide (docx)."
        return "", "Extension non supportée."
    except Exception as exc:  # noqa: BLE001
        return "", str(exc)


def save_improved_copy(
    pack_root: Path,
    source_rel_posix: str,
    improved_text: str,
    *,
    as_docx: bool,
) -> tuple[Optional[Path], Optional[Path], Optional[str]]:
    """
    Enregistre une nouvelle version sous _exports/ en conservant l'arborescence
    relative du fichier source (sans écraser l'original).

    Retourne (chemin .md, chemin .docx ou None, erreur).
    """
    root = Path(pack_root)
    if not root.is_dir():
        return None, None, "Dossier parcours introuvable."

    text = improved_text.strip()
    if not text:
        return None, None, "Texte vide."

    src = root / source_rel_posix.replace("/", os.sep)
    if not src.is_file():
        return None, None, "Fichier source introuvable pour l'export."

    rel = Path(source_rel_posix)
    parent_parts = rel.parent.parts
    stem = rel.stem
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    base_name = f"{stem}_amelioration_{ts}"

    out_dir = root / EXPORTS_DIRNAME
    for part in parent_parts:
        out_dir = out_dir / part
    out_dir.mkdir(parents=True, exist_ok=True)

    try:
        md_path = out_dir / f"{base_name}.md"
        md_path.write_text(text, encoding="utf-8")

        docx_path: Optional[Path] = None
        if as_docx:
            import docx

            docx_path = out_dir / f"{base_name}.docx"
            doc = docx.Document()
            for line in text.split("\n"):
                line = line.strip()
                if line:
                    doc.add_paragraph(line)
            doc.save(str(docx_path))

        return md_path, docx_path, None
    except Exception as exc:  # noqa: BLE001
        return None, None, str(exc)


def save_resource_sidecar(
    pack_root: Path, source_rel_posix: str, notes: str
) -> tuple[Optional[Path], Optional[str]]:
    """
    Enregistre liens YouTube / notes à côté de la leçon dans _exports/ (fichier .txt).
    """
    root = Path(pack_root)
    if not root.is_dir():
        return None, "Dossier parcours introuvable."
    raw = notes.strip()
    if not raw:
        return None, "Rien à enregistrer."

    rel = Path(source_rel_posix)
    parent_parts = rel.parent.parts
    stem = rel.stem
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")

    out_dir = root / EXPORTS_DIRNAME
    for part in parent_parts:
        out_dir = out_dir / part
    out_dir.mkdir(parents=True, exist_ok=True)

    path = out_dir / f"{stem}_ressources_{ts}.txt"
    try:
        path.write_text(raw, encoding="utf-8")
        return path, None
    except Exception as exc:  # noqa: BLE001
        return None, str(exc)


def format_preview_html(text: str, *, force_rtl: bool) -> str:
    """
    Texte échappé + bloc HTML pour iframe (darija / arabe + français mélangés).
    """
    return format_preview_html_lecture(text, force_rtl=force_rtl, variant="standard")


def format_preview_html_docx_rich(
    path: Path,
    *,
    force_rtl: bool,
) -> str | None:
    """
    Aperçu fidèle au Word : gras, italique, tableaux, RTL selon ``force_rtl``.
    Les couleurs de police du .docx ne sont pas reproduites, sauf la zone **Support visuel**
    (M0) mise en évidence comme dans le gabarit.
    """
    try:
        from ai.docx_preview_html import (
            build_docx_lecture_iframe_html,
            docx_preview_available,
        )
    except ImportError:
        return None
    p = Path(path)
    if not docx_preview_available(p):
        return None
    try:
        return build_docx_lecture_iframe_html(p, force_rtl=force_rtl)
    except Exception:  # noqa: BLE001
        return None


def format_preview_html_lecture(
    text: str, *, force_rtl: bool, variant: str = "lecture"
) -> str:
    """
    Prévisualisation pour lecture : police confortable, largeur max, fond neutre.
    ``variant`` : ``lecture`` (lisible) ou ``standard`` (légèrement plus compact).
    """
    safe = html.escape(text)
    direction = "rtl" if force_rtl else "auto"
    text_align = "right" if force_rtl else "start"
    if variant == "lecture":
        fs = "1.2rem"
        lh = "1.95"
        bg = "#ffffff"
        pad = "22px 28px"
        maxw = "720px"
    else:
        fs = "1.08rem"
        lh = "1.8"
        bg = "#fafafa"
        pad = "14px 16px"
        maxw = "100%"
    return f"""<!DOCTYPE html>
<html lang="ar"><head>
<meta charset="utf-8"/>
<link href="https://fonts.googleapis.com/css2?family=Noto+Naskh+Arabic:wght@400;700&family=Noto+Sans:wght@400;600&display=swap" rel="stylesheet">
<style>
  body {{ margin: 0; background: #e8eaed; }}
  .wrap {{ max-width: {maxw}; margin: 0 auto; background: {bg}; min-height: 100vh;
    box-shadow: 0 0 1px rgba(0,0,0,.08); }}
  .box {{
    direction: {direction};
    text-align: {text_align};
    unicode-bidi: plaintext;
    font-family: 'Noto Naskh Arabic', 'Noto Sans Arabic', 'Noto Sans', 'Segoe UI', sans-serif;
    font-size: {fs};
    line-height: {lh};
    white-space: pre-wrap;
    word-wrap: break-word;
    padding: {pad};
    min-height: 120px;
    color: #1a1a1a;
  }}
</style>
</head><body><div class="wrap"><div class="box">{safe}</div></div></body></html>"""
