"""
Conversion d’un .docx en HTML pour prévisualisation Streamlit (gras, italique, RTL).

Reproduit de façon proche l’affichage Word : paragraphes, runs (bold / italic / souligné),
sans couleur de police sauf le bloc **Support visuel** (repère prod, aligné sur le Word M0).
Tableaux simples (y compris dans les cellules).
"""
from __future__ import annotations

import difflib
import html
import re
from pathlib import Path
from typing import Union

from docx.document import Document as DocumentObject
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph

# Aligné sur ``format_module00_scripts._add_visual_support_block`` (titre exact)
_SUPPORT_VISUEL_TITLE = "Support visuel — fin de séquence (tournage / slide)"
_SEQ_TITLE_START = re.compile(r"(?i)^Séquence\s+\d+\s+—")


def _norm_block_text(s: str) -> str:
    return " ".join((s or "").split())


def _block_signature(block: Union[Paragraph, Table]) -> str:
    """Texte normalisé pour comparer deux .docx bloc à bloc (paragraphe ou tableau)."""
    if isinstance(block, Paragraph):
        return _norm_block_text(block.text or "")
    if isinstance(block, Table):
        parts: list[str] = []
        for row in block.rows:
            for cell in row.cells:
                parts.append(_norm_block_text(cell.text or ""))
        return "\u241e".join(parts)
    return ""


def compute_body_diff_marks_v1_vs_v2(
    doc_v1: DocumentObject,
    doc_v2: DocumentObject,
) -> dict[int, str]:
    """
    Indices des blocs du **corps** de ``doc_v2`` : ``add`` (nouveau) ou ``chg`` (remplace un bloc V1 différent).
    Même logique que l’aperçu Streamlit V2 vs V1.
    """
    s1 = [_block_signature(b) for b in _iter_block_items(doc_v1)]
    s2 = [_block_signature(b) for b in _iter_block_items(doc_v2)]
    sm = difflib.SequenceMatcher(None, s1, s2)
    mark: dict[int, str] = {}
    for tag, _i1, _i2, j1, j2 in sm.get_opcodes():
        if tag == "equal":
            continue
        if tag == "insert":
            for j in range(j1, j2):
                mark[j] = "add"
        elif tag == "replace":
            for j in range(j1, j2):
                mark[j] = "chg"
    return mark


def _v2_diff_marks_vs_v1(path_v1: Path, path_v2: Path) -> dict[int, str] | None:
    """
    Indices des blocs du **V2** : ``add`` (nouveau) ou ``chg`` (remplace un bloc V1 différent).
    Retourne ``None`` si comparaison impossible.
    """
    try:
        import docx
    except ImportError:
        return None
    p1, p2 = Path(path_v1), Path(path_v2)
    if not p1.is_file() or not p2.is_file():
        return None
    try:
        d1 = docx.Document(str(p1))
        d2 = docx.Document(str(p2))
    except Exception:  # noqa: BLE001
        return None
    return compute_body_diff_marks_v1_vs_v2(d1, d2)


def _iter_block_items(parent: Union[DocumentObject, _Cell]):
    """Parcourt les paragraphes et tableaux dans l’ordre du document (corps ou cellule)."""
    if isinstance(parent, DocumentObject):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise ValueError(f"Unsupported parent type: {type(parent)}")
    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


def _run_to_html(run) -> str:
    t = html.escape(run.text or "")
    if not t:
        return ""
    # Pas de couleur de police dans l’aperçu : lecture uniforme (noir / héritage CSS).
    style_parts: list[str] = []
    if run.bold:
        t = f"<strong>{t}</strong>"
    if run.italic:
        t = f"<em>{t}</em>"
    if run.underline:
        t = f"<u>{t}</u>"
    if style_parts:
        t = f'<span style="{";".join(style_parts)}">{t}</span>'
    return t


def _paragraph_to_html(para: Paragraph) -> str:
    if para.runs:
        inner = "".join(_run_to_html(r) for r in para.runs)
    else:
        inner = html.escape(para.text or "")
    if not inner.strip():
        return '<p class="docx-p docx-p-empty">&nbsp;</p>'
    return f'<p class="docx-p">{inner}</p>'


def _render_table(table: Table) -> str:
    rows_html: list[str] = []
    for row in table.rows:
        cells_html: list[str] = []
        for cell in row.cells:
            cell_inner: list[str] = []
            for block in _iter_block_items(cell):
                if isinstance(block, Paragraph):
                    cell_inner.append(_paragraph_to_html(block))
                elif isinstance(block, Table):
                    cell_inner.append(_render_table(block))
            cells_html.append(f'<td class="docx-td">{"".join(cell_inner)}</td>')
        rows_html.append(f"<tr>{''.join(cells_html)}</tr>")
    return f'<table class="docx-table"><tbody>{"".join(rows_html)}</tbody></table>'


def _block_to_inner_html(block: Union[Paragraph, Table]) -> str:
    if isinstance(block, Paragraph):
        return _paragraph_to_html(block)
    if isinstance(block, Table):
        return _render_table(block)
    return ""


def document_to_html_body(
    doc: DocumentObject, marks: dict[int, str] | None = None
) -> str:
    """
    Corps HTML. Si ``marks`` est fourni (indices → ``add`` | ``chg``), chaque bloc est enveloppé
    pour permettre le surlignage des ajouts / modifications par rapport à un autre .docx.
    Sans ``marks`` : enveloppe le bloc **Support visuel** M0 (titre + suite jusqu’à la prochaine séquence).
    """
    chunks: list[str] = []
    in_support = False
    for j, block in enumerate(_iter_block_items(doc)):
        if isinstance(block, Paragraph):
            txt = (block.text or "").strip()
            if _SEQ_TITLE_START.match(txt):
                in_support = False
        inner = _block_to_inner_html(block)
        if marks is not None:
            cls = "docx-block"
            m = marks.get(j)
            if m == "add":
                cls += " docx-diff-add"
            elif m == "chg":
                cls += " docx-diff-chg"
            chunks.append(f'<div class="{cls}">{inner}</div>')
        else:
            support_wrap = False
            if isinstance(block, Paragraph):
                t = (block.text or "").strip()
                if t == _SUPPORT_VISUEL_TITLE:
                    support_wrap = True
                    in_support = True
                elif in_support:
                    support_wrap = True
            elif isinstance(block, Table) and in_support:
                support_wrap = True
            if support_wrap:
                inner = f'<div class="docx-support-visuel">{inner}</div>'
            chunks.append(inner)
    return "\n".join(chunks)


def _iframe_shell(
    inner: str,
    *,
    force_rtl: bool,
    legend_html: str = "",
) -> str:
    direction = "rtl" if force_rtl else "auto"
    text_align = "right" if force_rtl else "start"
    fs = "1.2rem"
    lh = "1.95"
    bg = "#ffffff"
    pad = "22px 28px"
    maxw = "720px"

    # Contenu déjà échappé run par run ; pas de double escape
    return f"""<!DOCTYPE html>
<html lang="ar"><head>
<meta charset="utf-8"/>
<link href="https://fonts.googleapis.com/css2?family=Noto+Naskh+Arabic:wght@400;700&family=Noto+Sans:wght@400;600&display=swap" rel="stylesheet">
<style>
  body {{ margin: 0; background: #e8eaed; }}
  .wrap {{ max-width: {maxw}; margin: 0 auto; background: {bg}; min-height: 100vh;
    box-shadow: 0 0 1px rgba(0,0,0,.08); }}
  .box {{
    direction: {direction};
    text-align: {text_align};
    unicode-bidi: plaintext;
    font-family: 'Noto Naskh Arabic', 'Noto Sans Arabic', 'Noto Sans', 'Segoe UI', sans-serif;
    font-size: {fs};
    line-height: {lh};
    white-space: normal;
    word-wrap: break-word;
    padding: {pad};
    min-height: 120px;
    color: #1a1a1a;
  }}
  .docx-p {{ margin: 0.5em 0; }}
  .docx-p-empty {{ margin: 0.25em 0; min-height: 0.5em; }}
  .docx-table {{ border-collapse: collapse; width: 100%; margin: 0.75em 0; font-size: 0.95em; }}
  .docx-td {{ border: 1px solid #ddd; padding: 8px 10px; vertical-align: top; }}
  .docx-table .docx-p:first-child {{ margin-top: 0; }}
  .docx-table .docx-p:last-child {{ margin-bottom: 0; }}
  .docx-block {{ margin: 0; padding: 0; }}
  .docx-diff-add {{
    background: linear-gradient(90deg, rgba(40, 167, 69, 0.14) 0%, rgba(40, 167, 69, 0.06) 100%);
    border-inline-start: 4px solid #28a745;
    padding-inline-start: 10px;
    margin: 0.35em 0;
    border-radius: 2px;
  }}
  .docx-diff-chg {{
    background: linear-gradient(90deg, rgba(255, 193, 7, 0.18) 0%, rgba(255, 193, 7, 0.06) 100%);
    border-inline-start: 4px solid #e0a800;
    padding-inline-start: 10px;
    margin: 0.35em 0;
    border-radius: 2px;
  }}
  .diff-legend {{
    font-size: 0.82em;
    color: #444;
    margin-bottom: 14px;
    padding: 10px 12px;
    background: #f1f3f4;
    border-radius: 6px;
    line-height: 1.5;
    text-align: start;
    direction: ltr;
  }}
  .diff-legend span {{ display: inline-block; margin-inline-end: 14px; }}
  .lg-add {{ color: #1e7e34; font-weight: 600; }}
  .lg-chg {{ color: #856404; font-weight: 600; }}
  .docx-support-visuel {{
    color: #0b6e4f;
    margin: 0.35em 0;
    padding: 8px 10px;
    border-inline-start: 3px solid #0b6e4f;
    background: rgba(11, 110, 79, 0.06);
    border-radius: 2px;
  }}
  .docx-support-visuel .docx-p {{ margin: 0.35em 0; }}
  .docx-support-visuel strong {{ color: inherit; }}
</style>
</head><body><div class="wrap"><div class="box">
{legend_html}{inner}
</div></div></body></html>"""


def build_docx_lecture_iframe_html(path: Path, *, force_rtl: bool) -> str:
    """
    Document HTML complet (iframe) pour une leçon .docx, style lecture Streamlit.
    """
    try:
        import docx
    except ImportError:
        return ""

    p = Path(path)
    if not p.is_file() or p.suffix.lower() != ".docx":
        return ""

    doc = docx.Document(str(p))
    inner = document_to_html_body(doc)
    return _iframe_shell(inner, force_rtl=force_rtl)


def build_docx_lecture_iframe_html_diff(
    path_v2: Path, path_v1: Path, *, force_rtl: bool
) -> str:
    """
    Aperçu du **V2** avec surlignage des blocs **ajoutés** (vert) ou **modifiés** vs V1 (ambre).
    Bloc = paragraphe ou tableau, dans l’ordre du document.
    """
    try:
        import docx
    except ImportError:
        return ""

    p2 = Path(path_v2)
    p1 = Path(path_v1)
    if not p2.is_file() or p2.suffix.lower() != ".docx":
        return ""
    if not p1.is_file() or p1.suffix.lower() != ".docx":
        return build_docx_lecture_iframe_html(p2, force_rtl=force_rtl)

    try:
        doc = docx.Document(str(p2))
    except Exception:  # noqa: BLE001
        return ""

    marks = _v2_diff_marks_vs_v1(p1, p2)
    if marks is None or not marks:
        inner = document_to_html_body(doc)
        return _iframe_shell(inner, force_rtl=force_rtl)

    inner = document_to_html_body(doc, marks=marks)
    legend = """<div class="diff-legend" dir="ltr">
  <span class="lg-add">■ Vert</span> bloc nouveau par rapport à V1
  &nbsp;·&nbsp;
  <span class="lg-chg">■ Ambre</span> bloc modifié (remplace un bloc V1 différent)
</div>"""
    return _iframe_shell(inner, force_rtl=force_rtl, legend_html=legend)


def docx_preview_available(path: Path) -> bool:
    p = Path(path)
    return p.is_file() and p.suffix.lower() == ".docx"
