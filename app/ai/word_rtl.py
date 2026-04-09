"""
Direction de lecture droite → gauche (RTL) pour documents Word (.docx) en arabe / darija.

Word attend notamment l’élément OOXML ``w:bidi`` sur ``w:pPr`` (paragraphe) et,
pour une meilleure compatibilité, ``w:rtl`` sur ``w:rPr`` (portion de texte).

Usage :
    from ai.word_rtl import apply_rtl_to_document
    doc = Document(path)
    apply_rtl_to_document(doc)
    doc.save(path)
"""
from __future__ import annotations

from docx.document import Document as DocumentObject
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph


def _qn_bidi() -> str:
    return qn("w:bidi")


def _qn_rtl() -> str:
    return qn("w:rtl")


def _has_direct_child(parent, tag: str) -> bool:
    return any(child.tag == tag for child in parent)


def ensure_paragraph_rtl(paragraph: Paragraph) -> None:
    """Ajoute ``w:bidi`` sur le paragraphe si absent."""
    p_pr = paragraph._p.get_or_add_pPr()
    if not _has_direct_child(p_pr, _qn_bidi()):
        p_pr.append(OxmlElement("w:bidi"))


def ensure_run_rtl(run) -> None:
    """Ajoute ``w:rtl`` sur le run si absent (scripts complexes / arabe)."""
    r_pr = run._element.get_or_add_rPr()
    if not _has_direct_child(r_pr, _qn_rtl()):
        r_pr.append(OxmlElement("w:rtl"))


def _apply_paragraph_and_runs(paragraph: Paragraph) -> None:
    ensure_paragraph_rtl(paragraph)
    for run in paragraph.runs:
        ensure_run_rtl(run)


def _iter_paragraphs_in_cell(cell) -> None:
    for paragraph in cell.paragraphs:
        _apply_paragraph_and_runs(paragraph)
    for table in cell.tables:
        _apply_table_rtl(table)


def _apply_table_rtl(table: Table) -> None:
    for row in table.rows:
        for cell in row.cells:
            _iter_paragraphs_in_cell(cell)


def _apply_header_footer_part(part) -> None:
    if part is None:
        return
    for paragraph in part.paragraphs:
        _apply_paragraph_and_runs(paragraph)
    for table in part.tables:
        _apply_table_rtl(table)


def apply_rtl_to_document(document: DocumentObject, *, skip_body_paragraphs: int = 0) -> int:
    """
    Applique RTL à tous les paragraphes du document (corps, tableaux imbriqués,
    en-têtes et pieds de page de chaque section).

    ``skip_body_paragraphs`` : nombre de paragraphes du **corps** laissés en ordre
    par défaut (ex. 2 pour « Module » + « Leçon » en tête, bilingues FR/AR).

    Retourne le nombre de paragraphes du corps traités (hors ignorés).
    """
    count = 0
    for i, paragraph in enumerate(document.paragraphs):
        if i < skip_body_paragraphs:
            continue
        _apply_paragraph_and_runs(paragraph)
        count += 1
    for table in document.tables:
        _apply_table_rtl(table)
    for section in document.sections:
        _apply_header_footer_part(section.header)
        _apply_header_footer_part(section.footer)
        # Première page / pairs si utilisés
        if section.different_first_page_header_footer:
            _apply_header_footer_part(section.first_page_header)
            _apply_header_footer_part(section.first_page_footer)
    return count
