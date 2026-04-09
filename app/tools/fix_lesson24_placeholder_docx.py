"""Crée un .docx valide pour la leçon 24 si le fichier était vide (0 octet)."""
from __future__ import annotations

import sys
from pathlib import Path

import docx

_APP_DIR = Path(__file__).resolve().parent.parent
_REPO_ROOT = _APP_DIR.parent
if str(_APP_DIR) not in sys.path:
    sys.path.insert(0, str(_APP_DIR))

from ai.word_rtl import apply_rtl_to_document

TARGET = (
    _REPO_ROOT
    / "Input"
    / "Script"
    / "Parcours stage & emploi"
    / "02-Module 2- L'entreprise"
    / "24-Le fonctionnement d'une entreprise"
    / "Leçon 24 - Le fonctionnement d'une entreprise darija.docx"
)


def main() -> None:
    d = docx.Document()
    d.add_heading("24 — Le fonctionnement d'une entreprise (brouillon V1)", 0)
    d.add_paragraph(
        "[Placeholder Promind7] Ce fichier était vide. Remplace ce texte par le script oral "
        "(darija + termes techniques en français ou arabe standard)."
    )
    TARGET.parent.mkdir(parents=True, exist_ok=True)
    apply_rtl_to_document(d)
    d.save(str(TARGET))
    print("Enregistré :", TARGET)
    print("Taille :", TARGET.stat().st_size, "octets")


if __name__ == "__main__":
    main()
