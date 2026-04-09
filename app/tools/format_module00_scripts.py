"""
Régénère les 6 scripts Word du Module 0 (Études & mindset) au format proM7 :
- En-tête (gras noir, Calibri 11) : (1) ligne Module 0, (2) ligne Leçon — puis section Objectif
- Section الهدف : titre gras noir + un paragraphe en arabe uniquement
- Séquences : titres gras noirs « Séquence N — … » sans double numéro (N répété une seule fois)
- Dans une séquence, un paragraphe peut être un couple ``("phase1"|"phase2"|"phase3", texte)`` : titre de phase en gras (Leçon 1, sommaire)
- Marque : proM7 dans le texte latin
- Dans le corps des séquences : paires **…** (Markdown) → gras Word (comme M3–M7)
- Fin de chaque séquence : bloc « Support visuel — tournage » (FR) via ``module00_sequence_visuals.py`` — **titre + consignes** en couleur repère (séparé du script oral, noir) ; reste du document en noir.
- RTL : apply_rtl_to_document avant enregistrement

Le dossier parent est résolu par glob « 00-Module* » (tolère le nom « mindest » sur le disque).

Usage (depuis la racine V2) :
    python app/tools/format_module00_scripts.py
    python app/tools/format_module00_scripts.py --out-dir "C:\\chemin\\sortie"
    python app/tools/format_module00_scripts.py --lesson 2
    python app/tools/format_module00_scripts.py --lesson 2 --v2
      # écrit ``…-V2.docx`` (V1 = Word d’origine inchangé) ; texte en style normal (pas de surlignage V1/V2)
"""
from __future__ import annotations

import argparse
import re
import sys
from collections.abc import Callable, Sequence
from pathlib import Path

_APP_DIR = Path(__file__).resolve().parent.parent
_REPO_ROOT = _APP_DIR.parent
if str(_APP_DIR) not in sys.path:
    sys.path.insert(0, str(_APP_DIR))

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from ai.word_rtl import apply_rtl_to_document

FONT = "Calibri"
SIZE_PT = 11
# Bloc « Support visuel » uniquement (schéma / slide) — distingué du script darija
COLOR_SUPPORT_VISUEL = RGBColor(0x0B, 0x6E, 0x4F)

_PHASE_TAGS = frozenset({"phase1", "phase2", "phase3"})

SCRIPT_PACK = _REPO_ROOT / "Input" / "Script" / "Parcours stage & emploi"


def _find_module0_root() -> Path:
    if not SCRIPT_PACK.is_dir():
        raise FileNotFoundError(f"Introuvable : {SCRIPT_PACK}")
    candidates = sorted(SCRIPT_PACK.glob("00-Module*"))
    if not candidates:
        raise FileNotFoundError(f"Aucun dossier 00-Module* sous {SCRIPT_PACK}")
    return candidates[0]


def _set_run_font(run, *, bold: bool, color: RGBColor | None = None) -> None:
    run.bold = bold
    run.font.name = FONT
    run.font.size = Pt(SIZE_PT)
    if color is not None:
        run.font.color.rgb = color
    r = run._element
    rpr = r.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), FONT)
    rfonts.set(qn("w:hAnsi"), FONT)
    rfonts.set(qn("w:cs"), FONT)


def _add_bold_line(doc: Document, text: str, *, color: RGBColor | None = None) -> None:
    text = (text or "").strip()
    if not text:
        return
    p = doc.add_paragraph()
    r = p.add_run(text)
    _set_run_font(r, bold=True, color=color)


def _add_normal_paragraph(doc: Document, text: str) -> None:
    text = (text or "").strip()
    if not text:
        return
    p = doc.add_paragraph()
    r = p.add_run(text)
    _set_run_font(r, bold=False)


def _add_body_paragraph(doc: Document, text: str, *, color: RGBColor | None = None) -> None:
    """Corps de séquence : **…** → gras Word ; ``color`` optionnel (bloc schéma)."""
    text = (text or "").strip()
    if not text:
        return
    text = _normalize_brand(text)
    p = doc.add_paragraph()
    if "**" not in text:
        r = p.add_run(text)
        _set_run_font(r, bold=False, color=color)
        return
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        _set_run_font(r, bold=(i % 2 == 1), color=color)


def _add_colored_bold_markdown_paragraph(
    doc: Document, text: str, *, color: RGBColor | None = None
) -> None:
    """Ligne de phase (tout gras) ; **…** → gras ; ``color`` optionnel."""
    text = (text or "").strip()
    if not text:
        return
    text = _normalize_brand(text)
    p = doc.add_paragraph()
    if "**" not in text:
        r = p.add_run(text)
        _set_run_font(r, bold=True, color=color)
        return
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        _set_run_font(r, bold=True, color=color)


def _add_objectif_section(doc: Document, texte_ar: str) -> None:
    _add_bold_line(doc, "الهدف")
    _add_normal_paragraph(doc, texte_ar.strip())


def _add_sequence_title(doc: Document, title: str) -> None:
    t = (title or "").strip()
    if not t:
        return
    _add_bold_line(doc, t)


def _normalize_brand(text: str) -> str:
    t = text
    t = re.sub(r"\bProMind7\b", "proM7", t, flags=re.I)
    t = re.sub(r"\bPM7\b", "proM7", t)
    t = re.sub(r"\bP7\b", "proM7", t)
    return t


# Évite « Séquence 1 — 1. … » (double numéro) ; gère aussi l’espace avant le point en RTL (« 1 .عنوان »)
_SEQ_DEDUP = re.compile(
    r"(?i)Séquence\s+(\d+)\s*—\s*\1\s*(?:\.\s*|\s+\.\s+)(.+)",
)


def _normalize_sequence_title(title: str) -> str:
    """Un seul numéro dans le titre : Séquence N — <sujet> (sans répéter N.)."""
    t = (title or "").strip()
    if not t:
        return t
    m = _SEQ_DEDUP.match(t)
    if m:
        rest = m.group(2).strip()
        return f"Séquence {m.group(1)} — {rest}"
    return t


def _apply_normal_base_style(doc: Document) -> None:
    sty = doc.styles["Normal"]
    sty.font.name = FONT
    sty.font.size = Pt(SIZE_PT)


def _emit_sequence_body_paragraph(doc: Document, item: str | tuple[str, str]) -> None:
    """Corps de séquence : texte normal, ou (phase1|phase2|phase3, texte) en gras."""
    if isinstance(item, tuple) and len(item) == 2:
        tag, text = item[0], item[1]
        if tag in _PHASE_TAGS:
            _add_colored_bold_markdown_paragraph(doc, (text or "").strip(), color=None)
            return
        _add_body_paragraph(doc, (text or "").strip())
        return
    _add_body_paragraph(doc, str(item))


def _add_visual_support_block(doc: Document, text: str) -> None:
    """Bloc production : idées de slide / schéma (FR), pas destiné à l’oral darija tel quel."""
    text = (text or "").strip()
    if not text:
        return
    _add_bold_line(
        doc,
        "Support visuel — fin de séquence (tournage / slide)",
        color=COLOR_SUPPORT_VISUEL,
    )
    for line in text.split("\n"):
        line = line.strip()
        if line:
            _add_body_paragraph(doc, line, color=COLOR_SUPPORT_VISUEL)


def _seq(
    doc: Document,
    title: str,
    bodies: Sequence[str | tuple[str, str]],
    *,
    visuel_fin: str | None = None,
) -> None:
    _add_sequence_title(doc, _normalize_brand(_normalize_sequence_title(title)))
    for b in bodies:
        _emit_sequence_body_paragraph(doc, b)
    if visuel_fin:
        _add_visual_support_block(doc, visuel_fin)


_TOOLS_DIR = Path(__file__).resolve().parent
if str(_TOOLS_DIR) not in sys.path:
    sys.path.insert(0, str(_TOOLS_DIR))
from module00_rich_content import LESSONS  # noqa: E402
from module00_sequence_visuals import VISUEL_FIN_SEQUENCE  # noqa: E402


def _build_lesson(lesson_id: int) -> Document:
    """Contenu riche restauré depuis l’extract + complétions (voir module00_rich_content.py)."""
    doc = Document()
    _apply_normal_base_style(doc)
    data = LESSONS[lesson_id]
    cadre = data.get("cadre_lignes", ())
    if len(cadre) >= 2:
        _add_bold_line(doc, cadre[0].strip())
        _add_bold_line(doc, cadre[1].strip())
    else:
        # Secours si cadre incomplet
        _add_bold_line(doc, data.get("title", "").strip())
    _add_objectif_section(doc, data["objectif_ar"])
    for i, seq in enumerate(data["sequences"]):
        seq_title, bodies = seq[0], seq[1]
        vis_key = (lesson_id, i)
        vis_text = VISUEL_FIN_SEQUENCE.get(vis_key)
        _seq(doc, seq_title, bodies, visuel_fin=vis_text)
    return doc


def build_lesson_01() -> Document:
    return _build_lesson(1)


def build_lesson_02() -> Document:
    return _build_lesson(2)


def build_lesson_03() -> Document:
    return _build_lesson(3)


def build_lesson_04() -> Document:
    return _build_lesson(4)


def build_lesson_05() -> Document:
    return _build_lesson(5)


def build_lesson_06() -> Document:
    return _build_lesson(6)


# (dossier relatif, nom fichier, builder)
LESSON_TARGETS: list[tuple[str, str, Callable[[], Document]]] = [
    ("01-Introduction au Parcours stage et emploi", "Leçon 1 - Introduction au Parcours stage et emploi-darija.docx", build_lesson_01),
    ("02-Les études et le choix de la specialité", "Leçon 2 - Les études et le choix de la specialité-darija.docx", build_lesson_02),
    ("03-les lacunes et la difficulté technique", "Leçon 3 - les lacunes et la difficulté technique-darija.docx", build_lesson_03),
    ("04-la gestion du stress en tant qu'étudiant", "Leçon 4 - la gestion du stress en tant qu'étudiant-darija.docx", build_lesson_04),
    ("05-La gestion du temps et la discipline", "Leçon 5 - La gestion du temps et la discipline-darija.docx", build_lesson_05),
    ("06-le lien entre les études et le marché du travail", "Leçon 6 - le lien entre les études et le marché du travail-darija.docx", build_lesson_06),
]


def main() -> None:
    ap = argparse.ArgumentParser(description="Régénère les 6 scripts Word du Module 0 (études & mindset).")
    ap.add_argument(
        "--out-dir",
        type=str,
        default="",
        help="Dossier racine de sortie (défaut : module 00 détecté sous Input/Script/...).",
    )
    ap.add_argument(
        "--lesson",
        type=int,
        default=0,
        metavar="N",
        help="Ne régénérer que la leçon N (1–6). 0 = les six leçons.",
    )
    ap.add_argument(
        "--v2",
        action="store_true",
        help="Écrire {nom}-V2.docx ; V1 (fichier sans -V2) inchangé ; affichage normal (sans couleurs de diff vs V1).",
    )
    args = ap.parse_args()
    mod0 = _find_module0_root()
    out_root = Path(args.out_dir.strip()) if args.out_dir.strip() else mod0
    if args.out_dir.strip():
        out_root.mkdir(parents=True, exist_ok=True)

    targets = LESSON_TARGETS
    if int(args.lesson) != 0:
        n = int(args.lesson)
        if not (1 <= n <= len(LESSON_TARGETS)):
            print("Erreur : --lesson doit être entre 1 et", len(LESSON_TARGETS), file=sys.stderr)
            sys.exit(1)
        targets = [LESSON_TARGETS[n - 1]]

    for subdir, filename, builder in targets:
        if args.out_dir.strip():
            out_dir = out_root / subdir
        else:
            out_dir = mod0 / subdir
        out_dir.mkdir(parents=True, exist_ok=True)
        fp = Path(filename)
        if args.v2:
            out_path = out_dir / f"{fp.stem}-V2{fp.suffix}"
        else:
            out_path = out_dir / filename
        doc = builder()
        # Module + Leçon en tête : laisser ordre de lecture naturel (FR puis rappel AR entre parenthèses)
        apply_rtl_to_document(doc, skip_body_paragraphs=2)
        doc.save(str(out_path))
        print("OK", out_path)


if __name__ == "__main__":
    main()
