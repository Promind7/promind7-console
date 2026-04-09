# -*- coding: utf-8 -*-
"""
Régénère les scripts Word du Module 6 (outils numériques) à partir de :
- tools/module06_rich_content.py (Leçons 52–56)

Même gabarit que Module 3–5 : cadre → الهدف → Séquences ; **…** → gras Word.

**Sortie par défaut : V2** — écrit ``…-V2.docx`` ; le .docx **canonique** (sans ``-V2``) **n’est pas modifié**.
Texte en style normal dans le V2 (pas de surlignage de diff vs V1).

Usage :
  python app/tools/format_module06_scripts.py
  python app/tools/format_module06_scripts.py --lesson 52
  python app/tools/format_module06_scripts.py --in-place
"""
from __future__ import annotations

import argparse
import re
import sys
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

_APP_DIR = Path(__file__).resolve().parent.parent
_REPO_ROOT = _APP_DIR.parent
if str(_APP_DIR) not in sys.path:
    sys.path.insert(0, str(_APP_DIR))
_TOOLS = Path(__file__).resolve().parent
if str(_TOOLS) not in sys.path:
    sys.path.insert(0, str(_TOOLS))

from ai.word_rtl import apply_rtl_to_document
from module06_rich_content import LESSONS as LESSONS_M6

SCRIPT_PACK = _REPO_ROOT / "Input" / "Script" / "Parcours stage & emploi"

FONT = "Calibri"
SIZE_PT = 11

_SEQ_DEDUP = re.compile(
    r"(?i)Séquence\s+(\d+)\s*—\s*\1\s*(?:\.\s*|\s+\.\s+)(.+)",
)


def _find_module_root(glob_pat: str) -> Path:
    if not SCRIPT_PACK.is_dir():
        raise FileNotFoundError(f"Introuvable : {SCRIPT_PACK}")
    c = sorted(SCRIPT_PACK.glob(glob_pat))
    if not c:
        raise FileNotFoundError(f"Aucun dossier {glob_pat!r} sous {SCRIPT_PACK}")
    return c[0]


def _set_run_font(run, *, bold: bool, color: RGBColor | None = None) -> None:
    run.bold = bold
    run.font.name = FONT
    run.font.size = Pt(SIZE_PT)
    if color is not None:
        run.font.color.rgb = color
    r = run._element
    rpr = r.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), FONT)
    rfonts.set(qn("w:hAnsi"), FONT)
    rfonts.set(qn("w:cs"), FONT)


def _add_bold_line(doc: Document, text: str, *, color: RGBColor | None = None) -> None:
    text = (text or "").strip()
    if not text:
        return
    p = doc.add_paragraph()
    r = p.add_run(text)
    _set_run_font(r, bold=True, color=color)


def _add_normal_paragraph(doc: Document, text: str) -> None:
    text = (text or "").strip()
    if not text:
        return
    p = doc.add_paragraph()
    r = p.add_run(text)
    _set_run_font(r, bold=False)


def _add_body_paragraph(doc: Document, text: str) -> None:
    """Paragraphes des séquences : les paires **…** (style Markdown) deviennent du gras Word."""
    text = (text or "").strip()
    if not text:
        return
    text = _normalize_brand(text)
    p = doc.add_paragraph()
    if "**" not in text:
        r = p.add_run(text)
        _set_run_font(r, bold=False)
        return
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        _set_run_font(r, bold=(i % 2 == 1))


def _add_objectif_section(doc: Document, texte_ar: str) -> None:
    _add_bold_line(doc, "الهدف")
    _add_normal_paragraph(doc, texte_ar.strip())


def _normalize_brand(text: str) -> str:
    t = text
    t = re.sub(r"\bProMind7\b", "proM7", t, flags=re.I)
    t = re.sub(r"\bPM7\b", "proM7", t)
    t = re.sub(r"\bP7\b", "proM7", t)
    return t


def _normalize_sequence_title(title: str) -> str:
    t = (title or "").strip()
    if not t:
        return t
    m = _SEQ_DEDUP.match(t)
    if m:
        rest = m.group(2).strip()
        return f"Séquence {m.group(1)} — {rest}"
    return t


def _apply_normal_base_style(doc: Document) -> None:
    sty = doc.styles["Normal"]
    sty.font.name = FONT
    sty.font.size = Pt(SIZE_PT)


def _emit_sequence_body_paragraph(doc: Document, item: str | tuple[str, str]) -> None:
    if isinstance(item, tuple) and len(item) == 2:
        _add_body_paragraph(doc, (item[1] or "").strip())
        return
    _add_body_paragraph(doc, str(item))


def _seq(doc: Document, title: str, bodies: list[str] | tuple) -> None:
    _add_bold_line(doc, _normalize_brand(_normalize_sequence_title(title)))
    for b in bodies:
        _emit_sequence_body_paragraph(doc, b)


def _build_lesson(data: dict) -> Document:
    doc = Document()
    _apply_normal_base_style(doc)
    cadre = data.get("cadre_lignes", ())
    if len(cadre) >= 2:
        _add_bold_line(doc, cadre[0].strip())
        _add_bold_line(doc, cadre[1].strip())
    else:
        _add_bold_line(doc, data.get("title", "").strip())
    _add_objectif_section(doc, data["objectif_ar"])
    for seq_title, bodies in data["sequences"]:
        _seq(doc, seq_title, bodies)
    return doc


def _resolve_existing_docx(module_root: Path, lesson_id: int) -> Path:
    """Chemin du .docx canonique (de préférence **sans** suffixe ``-V2``)."""
    subdirs = sorted(module_root.glob(f"{lesson_id:02d}-*"))
    if not subdirs:
        raise FileNotFoundError(f"Aucun sous-dossier {lesson_id:02d}-* sous {module_root}")
    folder = subdirs[0]
    docxs = sorted(folder.glob("*.docx"))
    if not docxs:
        raise FileNotFoundError(f"Aucun .docx dans {folder}")
    non_v2 = [p for p in docxs if not p.stem.endswith("-V2")]
    return non_v2[0] if non_v2 else docxs[0]


def _out_path(base: Path, *, v2: bool) -> Path:
    if not v2:
        return base
    stem = base.stem
    if stem.endswith("-V2"):
        stem = stem[: -len("-V2")]
    return base.parent / f"{stem}-V2{base.suffix}"


def _write_module(
    module_glob: str,
    lessons: dict[int, dict],
    *,
    lesson_id: int | None,
    v2: bool,
) -> None:
    root = _find_module_root(module_glob)
    for lid in sorted(lessons):
        if lesson_id is not None and lid != lesson_id:
            continue
        base = _resolve_existing_docx(root, lid)
        out = _out_path(base, v2=v2)
        doc = _build_lesson(lessons[lid])
        apply_rtl_to_document(doc, skip_body_paragraphs=2)
        doc.save(str(out))
        print("OK", out)


def main() -> None:
    ap = argparse.ArgumentParser(description="Régénère les scripts Word module 6 (Python → .docx).")
    ap.add_argument(
        "--in-place",
        action="store_true",
        help="Écrase le .docx canonique (V1). Par défaut : uniquement -V2.docx.",
    )
    ap.add_argument(
        "--lesson",
        type=int,
        default=0,
        metavar="N",
        help="Ne régénérer que la leçon N (52–56). 0 = toutes.",
    )
    args = ap.parse_args()
    v2 = not args.in_place
    lesson_filter: int | None = None if args.lesson == 0 else args.lesson
    _write_module("06-Module*", LESSONS_M6, lesson_id=lesson_filter, v2=v2)
    msg = "Terminé module 6"
    if v2:
        msg += " (sortie -V2, V1 inchangé)"
    else:
        msg += " (écriture sur V1 — --in-place)"
    if lesson_filter is not None:
        msg += f" — leçon {lesson_filter}"
    msg += "."
    print(msg)


if __name__ == "__main__":
    main()
